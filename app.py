import streamlit as st
import pandas as pd
import plotly.express as px
import base64
import io
import time
from datetime import datetime
from github import Github, InputGitTreeElement, GithubException
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from streamlit_option_menu import option_menu
import matplotlib.pyplot as plt
from streamlit_drawable_canvas import st_canvas
from PIL import Image

# --- SESSION STATE INITIALIZATION ---
def init_session_state():
    """Initializes session state variables."""
    if 'data_loaded' not in st.session_state:
        st.session_state.data_loaded = False
    if 'rca_records' not in st.session_state:
        st.session_state.rca_records = []
    if 'capa_records' not in st.session_state:
        st.session_state.capa_records = []
    if 'car_counter' not in st.session_state:
        st.session_state.car_counter = 1
    if 'user' not in st.session_state:
        st.session_state.user = "Anonymous User"
    if 'pareto_items' not in st.session_state:
        st.session_state.pareto_items = [{"cause": "", "frequency": 1}]
    if 'rca_record_type' not in st.session_state:
        st.session_state.rca_record_type = "Internal"

# --- GITHUB INTEGRATION ---
def get_github_repo():
    """Connects to the GitHub repository using the token from secrets."""
    try:
        g = Github(st.secrets["GITHUB_TOKEN"])
        repo = g.get_repo(st.secrets["GITHUB_REPO"])
        return repo
    except Exception as e:
        st.error(f"Error connecting to GitHub: {e}")
        return None

def load_data_from_github():
    """Loads RCA and CAPA data from GitHub repository."""
    repo = get_github_repo()
    if repo:
        # Load RCA data
        try:
            contents = repo.get_contents("rca_data.csv")
            rca_data = base64.b64decode(contents.content).decode('utf-8')
            st.session_state.rca_records = pd.read_csv(io.StringIO(rca_data)).to_dict('records')
        except GithubException as e:
            if e.status == 404:
                st.warning("'rca_data.csv' not found. Starting fresh.")
            else:
                st.error(f"Error loading rca_data.csv: {e}")
            st.session_state.rca_records = []
        except Exception as e:
            st.error(f"Unexpected error loading rca_data.csv: {e}")
            st.session_state.rca_records = []

        # Load CAPA data
        try:
            contents = repo.get_contents("capa_data.csv")
            capa_data = base64.b64decode(contents.content).decode('utf-8')
            st.session_state.capa_records = pd.read_csv(io.StringIO(capa_data)).to_dict('records')
            # Find the highest car_number for the current year
            if st.session_state.capa_records:
                current_year = datetime.now().year
                max_car = max(
                    [int(rec.get('car_number', '').split('-')[-1]) for rec in st.session_state.capa_records if rec.get('car_number', '').startswith(f"CAR-{current_year}-")],
                    default=0
                )
                st.session_state.car_counter = max_car + 1
            else:
                st.session_state.car_counter = 1
        except GithubException as e:
            if e.status == 404:
                st.warning("'capa_data.csv' not found. Starting fresh.")
            else:
                st.error(f"Error loading capa_data.csv: {e}")
            st.session_state.capa_records = []
            st.session_state.car_counter = 1
        except Exception as e:
            st.error(f"Unexpected error loading capa_data.csv: {e}")
            st.session_state.capa_records = []
            st.session_state.car_counter = 1
    else:
        st.warning("GitHub repository not connected. Cannot load data.")
        st.session_state.rca_records = []
        st.session_state.capa_records = []
        st.session_state.car_counter = 1
    st.session_state.data_loaded = True

def save_data_to_github():
    """Saves RCA and CAPA dataframes as CSVs to the GitHub repository."""
    repo = get_github_repo()
    if not repo:
        st.error("Cannot save: GitHub repository not connected.")
        return False
    try:
        rca_df = pd.DataFrame(st.session_state.rca_records)
        capa_df = pd.DataFrame(st.session_state.capa_records)
        files_to_commit = [
            InputGitTreeElement("rca_data.csv", "100644", "blob", rca_df.to_csv(index=False)),
            InputGitTreeElement("capa_data.csv", "100644", "blob", capa_df.to_csv(index=False))
        ]
        master_ref = repo.get_git_ref("heads/main")
        base_tree = repo.get_git_tree(master_ref.object.sha)
        tree = repo.create_git_tree(files_to_commit, base_tree)
        commit_message = f"QMS data update: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        parent = repo.get_git_commit(master_ref.object.sha)
        commit = repo.create_git_commit(commit_message, tree, [parent])
        master_ref.edit(commit.sha)
        st.success("Data successfully saved to GitHub!")
        return True
    except Exception as e:
        st.error(f"Error saving to GitHub: {e}")
        return False

# --- DOCX REPORT GENERATOR ---
def generate_docx(rca_data, capa_data):
    """Generates a professional DOCX report with Brafe branding, logo, and signatures."""
    try:
        doc = Document()
        # Add Brafe logo at the top
        try:
            logo = doc.add_picture("brafe-logo.png", width=Inches(2))
            logo_paragraph = doc.paragraphs[-1]
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Spacer after logo
        except FileNotFoundError:
            p = doc.add_paragraph("Brafe Engineering Logo Not Found")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red for visibility
        # Main title
        title = doc.add_heading('Corrective Action Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(16)
        title.runs[0].font.color.rgb = RGBColor(0, 40, 85)  # Brafe Blue
        doc.add_paragraph()  # Spacer
        # Helper function to add section heading
        def add_section_heading(text, level=1):
            heading = doc.add_heading(text, level=level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading.runs[0].font.color.rgb = RGBColor(255, 255, 255)  # White text
            heading.runs[0].font.size = Pt(14)
            paragraph_format = heading.paragraph_format
            paragraph_format.space_before = Pt(12)
            paragraph_format.space_after = Pt(6)
            # Add background color
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '002855')  # Brafe Blue (hex)
            heading._element.get_or_add_pPr().append(shading)
        # Helper function to add key-value pair
        def add_key_value_pair(key, value, use_accent=False):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run_key = p.add_run(f"{key}: ")
            run_key.bold = True
            run_key.font.size = Pt(11)
            run_key.font.color.rgb = RGBColor(0, 102, 204) if use_accent else RGBColor(0, 0, 0)  # Brafe Light Blue or Black
            run_value = p.add_run(str(value))
            run_value.font.size = Pt(11)
            # Add light gray shading for contrast
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F0F0F0')  # Light Gray
            p._element.get_or_add_pPr().append(shading)
        # Helper function to add body text
        def add_body_text(text):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        # Helper function to add horizontal line
        def add_horizontal_line():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run()
            run.add_break(WD_BREAK.LINE)
            border = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')  # Border thickness
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '002855')  # Brafe Blue
            border.append(bottom)
            p._element.get_or_add_pPr().append(border)
        # Section 1: Report Details
        add_section_heading('1. Report Identification')
        add_body_text("This section provides essential details for the identification and traceability of the Corrective Action Report.")
        add_key_value_pair('Report Number', capa_data.get('car_number', 'N/A'), use_accent=True)
        add_key_value_pair('Prepared By', rca_data.get('generated_by', 'N/A'))
        add_key_value_pair('Date Prepared', datetime.now().strftime('%Y-%m-%d'), use_accent=True)
        add_key_value_pair('Record Type', rca_data.get('record_type', 'N/A'))
        if rca_data.get('record_type') == 'Customer':
            add_key_value_pair('Customer Name', rca_data.get('customer_name', 'N/A'), use_accent=True)
            add_key_value_pair('Purchase Order Number', rca_data.get('po_number', 'N/A'))
            add_key_value_pair('Work Order Number', rca_data.get('work_order', 'N/A'), use_accent=True)
        add_horizontal_line()
        # Section 2: Non-Conformance Description
        add_section_heading('2. Non-Conformance Description')
        add_body_text("This section describes the observed non-conformance, including details of the issue and its context.")
        add_body_text(rca_data.get('problem_description', 'No description provided.'))
        add_horizontal_line()
        # Section 3: Root Cause Analysis
        add_section_heading(f"3. Root Cause Analysis ({rca_data.get('rca_technique', 'N/A')})")
        add_body_text(f"This section outlines the root cause analysis performed using the {rca_data.get('rca_technique', 'N/A')} methodology to identify the underlying causes of the non-conformance.")
        technique = rca_data.get('rca_technique')
        details = rca_data.get('technique_details', {})
        if technique == '5 Whys':
            for i, why in enumerate(details.get('whys', []), 1):
                add_key_value_pair(f"Why {i}", why, use_accent=(i % 2 == 0))
        elif technique == 'Fishbone Diagram':
            for category, causes in details.get('fishbone', {}).items():
                if causes:
                    add_key_value_pair(category, ', '.join(causes), use_accent=True)
        elif technique == 'Pareto Analysis':
            add_body_text("Pareto Analysis identifies the most significant contributing factors to the non-conformance, prioritizing corrective actions based on frequency and impact.")
            if 'pareto_chart' in rca_data and isinstance(rca_data['pareto_chart'], io.BytesIO):
                rca_data['pareto_chart'].seek(0)
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.add_run().add_picture(rca_data['pareto_chart'], width=Inches(6))
        add_horizontal_line()
        # Section 4: Corrective and Preventive Actions
        add_section_heading('4. Corrective and Preventive Actions')
        add_body_text("This section details the corrective actions to address the root cause and preventive actions to mitigate future occurrences.")
        add_key_value_pair('Corrective Action', capa_data.get('corrective_action', 'No action specified.'), use_accent=True)
        add_key_value_pair('Preventive Action', capa_data.get('preventive_action', 'No action specified.'))
        add_key_value_pair('Responsible Person/Team', capa_data.get('responsible', 'N/A'), use_accent=True)
        add_key_value_pair('Due Date', capa_data.get('due_date', 'N/A'))
        add_key_value_pair('Status', capa_data.get('status', 'N/A'), use_accent=True)
        add_horizontal_line()
        # Section 5: Approvals
        add_section_heading('5. Approvals')
        add_body_text("This section includes signatures from the responsible person and quality assurance approver, confirming review and approval of the report.")
        responsible_sig = capa_data.get('responsible_sig')
        approver_sig = capa_data.get('approver_sig')
        if responsible_sig and responsible_sig.startswith('data:image/png;base64,'):
            try:
                img_bytes = base64.b64decode(responsible_sig.split(',')[1])
                img_stream = io.BytesIO(img_bytes)
                p = doc.add_paragraph()
                p.add_run('Responsible Person Signature: ')
                p.add_run().add_picture(img_stream, width=Inches(2))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except Exception:
                add_key_value_pair('Responsible Person Signature', '[Signature Not Available]', use_accent=True)
        else:
            add_key_value_pair('Responsible Person Signature', '[Signature Not Available]', use_accent=True)
        if approver_sig and approver_sig.startswith('data:image/png;base64,'):
            try:
                img_bytes = base64.b64decode(approver_sig.split(',')[1])
                img_stream = io.BytesIO(img_bytes)
                p = doc.add_paragraph()
                p.add_run('QA Approver Signature: ')
                p.add_run().add_picture(img_stream, width=Inches(2))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except Exception:
                add_key_value_pair('QA Approver Signature', '[Signature Not Available]', use_accent=True)
        else:
            add_key_value_pair('QA Approver Signature', '[Signature Not Available]', use_accent=True)
        add_key_value_pair('Approval Date', datetime.now().strftime('%Y-%m-%d'), use_accent=True)
        add_horizontal_line()
        # Section 6: Supporting Evidence
        if rca_data.get('images'):
            add_section_heading('6. Supporting Evidence')
            add_body_text("This section includes photographic or visual evidence supporting the non-conformance investigation.")
            for img_data in rca_data['images']:
                try:
                    img_bytes = base64.b64decode(img_data.split(",")[1])
                    img_stream = io.BytesIO(img_bytes)
                    img_p = doc.add_paragraph()
                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_p.add_run().add_picture(img_stream, width=Inches(6))
                except Exception as e:
                    add_body_text(f"Error processing evidence image: {e}")
            add_horizontal_line()
        # Save to BytesIO
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream.getvalue()
    except Exception as e:
        error_message = f"An unexpected error occurred during DOCX generation: {e}"
        print(error_message)
        return {"error": error_message}

# --- UI RENDERING FUNCTIONS ---
def render_dashboard():
    """Renders the main dashboard page."""
    st.title("📊 QMS Dashboard")
    rca_df = pd.DataFrame(st.session_state.rca_records)
    capa_df = pd.DataFrame(st.session_state.capa_records)
    if rca_df.empty:
        st.info("No data available. Create your first RCA record to get started.")
        return
    st.subheader("Quality Management System")
    st.markdown("Track, analyze, and improve quality processes at Brafe Engineering")
    c1, c2, c3 = st.columns(3)
    open_capas = capa_df[capa_df['status'] != 'Closed'] if not capa_df.empty else pd.DataFrame()
    c1.metric("Total RCA Records", len(rca_df))
    c2.metric("Total CAPA Records", len(capa_df))
    c3.metric("Open CAPAs", len(open_capas))
    st.divider()
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("CAPA Status Distribution")
        if not capa_df.empty:
            status_counts = capa_df['status'].value_counts()
            fig = px.pie(status_counts, values=status_counts.values, names=status_counts.index, hole=.3,
                         color_discrete_map={'Open': '#EF553B', 'In Progress': '#FF97FF', 'Completed': '#00CC96', 'Closed': '#636EFA'})
            fig.update_layout(showlegend=True, margin=dict(t=0, b=0, l=0, r=0))
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.text("No CAPA data.")
    with c2:
        st.subheader("RCA Record Types")
        type_counts = rca_df['record_type'].value_counts()
        fig = px.bar(type_counts, x=type_counts.index, y=type_counts.values, labels={'x': 'Record Type', 'y': 'Count'})
        fig.update_layout(showlegend=False, margin=dict(t=0, b=0, l=0, r=0))
        st.plotly_chart(fig, use_container_width=True)
    st.divider()
    st.subheader("Recent Activity Log")
    if not capa_df.empty and not rca_df.empty:
        merged_df = pd.merge(rca_df, capa_df, left_on='id', right_on='rca_id', how='left', suffixes=('_rca', '_capa'))
        st.dataframe(merged_df[['id_rca', 'problem_description', 'status', 'due_date']].tail(), use_container_width=True)
    else:
        st.dataframe(rca_df[['id', 'problem_description', 'record_type']].tail(), use_container_width=True)

def render_create_rca():
    """Renders the page for creating a new RCA record."""
    st.title("📝 Create Root Cause Analysis")
    st.subheader("Root Cause Analysis")
    st.markdown("Identify the fundamental causes of quality issues")
    # Record type selection outside form to persist state
    record_type = st.radio(
        "Record Type",
        ["Internal", "Customer"],
        index=0 if st.session_state.rca_record_type == "Internal" else 1,
        key="rca_record_type_radio"
    )
    st.session_state.rca_record_type = record_type
    rca_techniques = {
        "5 Whys": "A simple, iterative technique to explore the cause-and-effect relationships underlying a problem.",
        "Fishbone Diagram": "Helps visualize potential causes of a problem by grouping them into major categories.",
        "Pareto Analysis": "Identifies the most significant factors using the 80/20 rule to prioritize efforts."
    }
    with st.form("rca_form_main", clear_on_submit=True):
        st.subheader("1. General Information")
        c1, c2 = st.columns(2)
        with c1:
            customer_name = st.text_input(
                "Customer Name",
                disabled=(st.session_state.rca_record_type == "Internal"),
                key="customer_name"
            )
        with c2:
            po_number = st.text_input(
                "Purchase Order (PO)",
                disabled=(st.session_state.rca_record_type == "Internal"),
                key="po_number"
            )
            work_order = st.text_input(
                "Work Order",
                disabled=(st.session_state.rca_record_type == "Internal"),
                key="work_order"
            )
        st.subheader("2. Problem Details")
        problem_description = st.text_area(
            "Problem Description",
            height=100,
            placeholder="Describe the issue, what happened, and where it was observed.",
            key="problem_description"
        )
        st.subheader("3. Root Cause Analysis Technique")
        technique = st.selectbox("Select RCA Technique", options=list(rca_techniques.keys()), key="technique")
        with st.expander("What is this technique?"):
            st.info(rca_techniques[technique])
        technique_details = {}
        if technique == '5 Whys':
            st.markdown("**Answer the 'Why?' questions to drill down to the root cause:**")
            whys = [st.text_input(f"Why {i+1}?", key=f"why{i}") for i in range(5)]
            technique_details['whys'] = [w for w in whys if w]
        elif technique == 'Fishbone Diagram':
            fishbone_data = {}
            categories = ['Manpower', 'Method', 'Machine', 'Material', 'Measurement', 'Environment']
            st.markdown("**Enter potential causes for each category:**")
            cols = st.columns(2)
            for i, cat in enumerate(categories):
                with cols[i % 2]:
                    causes = st.text_area(f"{cat} Causes", height=80, key=f"fishbone_{cat}")
                    fishbone_data[cat] = [c.strip() for c in causes.split('\n') if c.strip()]
            technique_details['fishbone'] = fishbone_data
        elif technique == 'Pareto Analysis':
            st.markdown("**Enter defect types and their frequency:**")
            if 'pareto_items' not in st.session_state:
                st.session_state.pareto_items = [{"cause": "", "frequency": 1}]
            pareto_container = st.container()
            for i in range(len(st.session_state.pareto_items)):
                with pareto_container:
                    c1, c2, c3 = st.columns([4, 2, 1])
                    with c1:
                        st.session_state.pareto_items[i]["cause"] = st.text_input("Cause", value=st.session_state.pareto_items[i]["cause"], key=f"cause_{i}")
                    with c2:
                        st.session_state.pareto_items[i]["frequency"] = st.number_input("Frequency", value=st.session_state.pareto_items[i]["frequency"], min_value=1, key=f"freq_{i}")
                    with c3:
                        st.write("")  # Empty space for alignment
            technique_details['pareto'] = [item for item in st.session_state.pareto_items if item['cause']]
        st.subheader("4. Evidence")
        images = st.file_uploader("Upload Evidence Photos", type=["jpg", "jpeg", "png"], accept_multiple_files=True, key="evidence_photos")
        submitted = st.form_submit_button("✅ Save RCA Record", use_container_width=True)
        if submitted:
            if not problem_description:
                st.error("Problem Description is required!")
            else:
                image_data = []
                for img in images:
                    img_bytes = img.getvalue()
                    img_base64 = base64.b64encode(img_bytes).decode()
                    image_data.append(f"data:{img.type};base64,{img_base64}")
                rca_record = {
                    "id": f"RCA-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
                    "record_type": st.session_state.rca_record_type,
                    "customer_name": customer_name if st.session_state.rca_record_type == "Customer" else "",
                    "po_number": po_number if st.session_state.rca_record_type == "Customer" else "",
                    "work_order": work_order if st.session_state.rca_record_type == "Customer" else "",
                    "problem_description": problem_description,
                    "rca_technique": technique,
                    "technique_details": technique_details,
                    "generated_by": st.session_state.user,
                    "created_at": datetime.now().isoformat(),
                    "images": image_data
                }
                if technique == 'Pareto Analysis' and technique_details.get('pareto'):
                    df = pd.DataFrame(technique_details['pareto']).sort_values('frequency', ascending=False)
                    df['cumulative_percentage'] = (df['frequency'].cumsum() / df['frequency'].sum()) * 100
                    fig, ax1 = plt.subplots()
                    ax1.bar(df['cause'], df['frequency'], color='C0')
                    ax1.set_ylabel('Frequency', color='C0')
                    ax1.tick_params(axis='y', labelcolor='C0')
                    ax1.tick_params(axis='x', rotation=45)
                    ax2 = ax1.twinx()
                    ax2.plot(df['cause'], df['cumulative_percentage'], color='C1', marker='o', ms=5)
                    ax2.set_ylabel('Cumulative Percentage', color='C1')
                    ax2.tick_params(axis='y', labelcolor='C1')
                    ax2.set_ylim([0, 110])
                    plt.title('Pareto Analysis')
                    fig.tight_layout()
                    buf = io.BytesIO()
                    fig.savefig(buf, format='png')
                    plt.close(fig)
                    rca_record['pareto_chart'] = buf
                st.session_state.rca_records.append(rca_record)
                with st.spinner("Saving RCA record..."):
                    time.sleep(1)
                st.success("RCA record created successfully!")
                st.balloons()
                st.session_state.pareto_items = [{"cause": "", "frequency": 1}]
                st.rerun()
    st.markdown("**Manage Causes:**")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("➕ Add Cause", key="add_cause"):
            st.session_state.pareto_items.append({"cause": "", "frequency": 1})
            st.rerun()
    with c2:
        if st.button("🗑️ Remove Last Cause", key="remove_last_cause"):
            if len(st.session_state.pareto_items) > 1:
                st.session_state.pareto_items.pop()
                st.rerun()

def render_create_capa():
    """Renders the page for creating a new CAPA record."""
    st.title("🛡️ Create Corrective/Preventive Action (CAPA)")
    st.subheader("Corrective and Preventive Actions")
    st.markdown("Implement solutions to address root causes and prevent recurrence of non-conformances")
    if not st.session_state.rca_records:
        st.warning("No RCA records found. Please create an RCA first.")
        return
    rca_options = {rca['id']: f"{rca['id']} - {rca.get('problem_description', 'N/A')[:50]}..." for rca in st.session_state.rca_records}
    selected_rca_id = st.selectbox("Select the RCA record to address:", options=list(rca_options.keys()), format_func=lambda x: rca_options[x], key="rca_selection")
    rca_data = next((r for r in st.session_state.rca_records if r['id'] == selected_rca_id), None)
    if rca_data:
        with st.expander("Selected RCA Details", expanded=True):
            st.write(f"**Non-Conformance:** {rca_data['problem_description']}")
            st.write(f"**RCA Technique:** {rca_data.get('rca_technique', 'N/A')}")
        with st.form("capa_form", clear_on_submit=True):
            st.subheader("1. Action Plan")
            car_number = st.text_input("Report Number", value=f"CAR-{datetime.now().year}-{st.session_state.car_counter:03d}", disabled=True, key="car_number")
            c1, c2 = st.columns(2)
            with c1:
                corrective_action = st.text_area("Corrective Action", height=150, help="Actions to eliminate the cause of the detected non-conformance.", key="corrective_action")
            with c2:
                preventive_action = st.text_area("Preventive Action", height=150, help="Actions to prevent recurrence of the non-conformance.", key="preventive_action")
            st.subheader("2. Assignment and Status")
            c1, c2, c3 = st.columns(3)
            with c1:
                responsible = st.text_input("Responsible Person/Team", key="responsible_person")
            with c2:
                due_date = st.date_input("Due Date", key="due_date")
            with c3:
                status = st.selectbox("Status", ["Open", "In Progress", "Completed", "Closed"], key="capa_status")
            st.subheader("3. Signatures")
            c1, c2 = st.columns(2)
            with c1:
                st.write("Responsible Person's Signature")
                responsible_sig = st_canvas(
                    fill_color="rgba(255, 165, 0, 0.3)",
                    stroke_width=2,
                    stroke_color="#000000",
                    background_color="#FFFFFF",
                    height=150,
                    key="responsible_sig_canvas"
                )
            with c2:
                st.write("QA Approver's Signature")
                approver_sig = st_canvas(
                    fill_color="rgba(255, 165, 0, 0.3)",
                    stroke_width=2,
                    stroke_color="#000000",
                    background_color="#FFFFFF",
                    height=150,
                    key="approver_sig_canvas"
                )
            submitted = st.form_submit_button("✅ Save CAPA Record", use_container_width=True)
            if submitted:
                if not corrective_action or not preventive_action or not responsible:
                    st.error("Corrective Action, Preventive Action, and Responsible Person are required!")
                else:
                    # Process signatures as images
                    responsible_sig_data = None
                    if hasattr(responsible_sig, 'image_data') and responsible_sig.image_data is not None:
                        try:
                            img = Image.fromarray(responsible_sig.image_data)
                            buf = io.BytesIO()
                            img.save(buf, format='PNG')
                            responsible_sig_data = f"data:image/png;base64,{base64.b64encode(buf.getvalue()).decode()}"
                        except Exception:
                            responsible_sig_data = None
                    approver_sig_data = None
                    if hasattr(approver_sig, 'image_data') and approver_sig.image_data is not None:
                        try:
                            img = Image.fromarray(approver_sig.image_data)
                            buf = io.BytesIO()
                            img.save(buf, format='PNG')
                            approver_sig_data = f"data:image/png;base64,{base64.b64encode(buf.getvalue()).decode()}"
                        except Exception:
                            approver_sig_data = None
                    capa_record = {
                        "id": f"CAPA-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
                        "rca_id": selected_rca_id,
                        "car_number": car_number,
                        "corrective_action": corrective_action,
                        "preventive_action": preventive_action,
                        "responsible": responsible,
                        "due_date": due_date.isoformat(),
                        "status": status,
                        "responsible_sig": responsible_sig_data,
                        "approver_sig": approver_sig_data,
                        "created_at": datetime.now().isoformat()
                    }
                    st.session_state.capa_records.append(capa_record)
                    st.session_state.car_counter += 1
                    with st.spinner("Saving CAPA record..."):
                        time.sleep(1)
                    st.success("CAPA record created successfully!")
                    st.balloons()
                    st.rerun()

def render_generate_report():
    """Renders the page for generating a DOCX report."""
    st.title("📄 Generate Corrective Action Report")
    st.subheader("Generate Professional Reports")
    st.markdown("Create formalized Corrective Action Reports for quality management")
    if not st.session_state.capa_records:
        st.warning("No CAPA records found. Create a CAPA first.")
        return
    capa_options = {capa['car_number']: f"{capa['car_number']} - RCA: {capa['rca_id']}" for capa in st.session_state.capa_records}
    selected_car_number = st.selectbox("Select a report to generate:", options=list(capa_options.keys()), format_func=lambda x: capa_options[x], key="car_selection")
    capa_data = next((c for c in st.session_state.capa_records if c.get('car_number') == selected_car_number), None)
    if capa_data:
        rca_data = next((r for r in st.session_state.rca_records if r['id'] == capa_data['rca_id']), None)
        if not rca_data:
            st.error("Linked RCA record not found! Cannot generate report.")
            return
        st.subheader(f"Preview for {selected_car_number}")
        if st.button("🚀 Generate Report", use_container_width=True, key="generate_docx"):
            with st.spinner("Creating your report..."):
                progress_bar = st.progress(0)
                for percent in range(0, 101, 10):
                    time.sleep(0.1)
                    progress_bar.progress(percent)
                docx_output = generate_docx(rca_data, capa_data)
                if isinstance(docx_output, dict) and "error" in docx_output:
                    st.error(f"Failed to generate report: {docx_output['error']}")
                else:
                    st.success("Report generated successfully!")
                    st.balloons()
                    st.download_button(
                        label="📥 Download Report",
                        data=docx_output,
                        file_name=f"{capa_data['car_number']}_Corrective_Action_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="download_docx"
                    )
    else:
        st.warning("Please select a valid report.")

def render_settings():
    """Renders the settings page."""
    st.title("⚙️ System Settings")
    st.subheader("System Configuration")
    st.markdown("Manage your preferences and system data")
    st.subheader("User Preferences")
    username = st.text_input("Your Name (for 'Prepared By' field)", value=st.session_state.get('user', ''), key="username")
    if st.button("Save User Preference", key="save_user"):
        st.session_state.user = username
        st.success("Username updated!")
    st.divider()
    st.subheader("System Data Management")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("💾 Save All Data to GitHub", use_container_width=True, key="save_github"):
            with st.spinner("Saving to GitHub..."):
                save_data_to_github()
                time.sleep(1)
            st.success("Data saved successfully!")
    with c2:
        if st.button("⚠️ Reset All Local Data", type="primary", use_container_width=True, key="reset_data"):
            if st.checkbox("Confirm data reset", key="confirm_reset"):
                for key in ['rca_records', 'capa_records', 'car_counter', 'data_loaded', 'pareto_items', 'rca_record_type']:
                    if key in st.session_state:
                        del st.session_state[key]
                init_session_state()
                st.success("All local session data has been reset!")
                st.rerun()
    st.info(f"""
        **Current Session Info:**
        - **RCA Records:** {len(st.session_state.rca_records)}
        - **CAPA Records:** {len(st.session_state.capa_records)}
        - **Next Report Number:** CAR-{datetime.now().year}-{st.session_state.car_counter:03d}
        - **GitHub Connection:** {'Active' if get_github_repo() else 'Inactive'}
    """)

# --- MAIN APP LOGIC ---
def main():
    """Main function to run the Streamlit app."""
    init_session_state()
    if "GITHUB_TOKEN" in st.secrets and not st.session_state.data_loaded:
        with st.spinner("Loading data from repository..."):
            progress_bar = st.progress(0)
            status_text = st.empty()
            for i in range(1, 11):
                progress_bar.progress(i * 10)
                status_text.text(f"Loading data... {i * 10}%")
                time.sleep(0.1)
            load_data_from_github()
            progress_bar.empty()
            status_text.empty()
    with st.sidebar:
        try:
            st.image("brafe-logo.png")
        except Exception:
            st.warning("Brafe logo not found.")
        st.markdown(f"Welcome, **{st.session_state.user}**")
        selected = option_menu(
            menu_title="Main Menu",
            options=["Dashboard", "Create RCA", "Create CAPA", "Generate Report", "Settings"],
            icons=["speedometer2", "journal-plus", "shield-check", "file-earmark-word", "gear-wide-connected"],
            menu_icon="cast",
            default_index=0,
            key="main_menu"
        )
        st.sidebar.divider()
        st.sidebar.info("This QMS application is for internal use at Brafe Engineering.")
    page_map = {
        "Dashboard": render_dashboard,
        "Create RCA": render_create_rca,
        "Create CAPA": render_create_capa,
        "Generate Report": render_generate_report,
        "Settings": render_settings
    }
    page_map[selected]()

if __name__ == "__main__":
    main()
